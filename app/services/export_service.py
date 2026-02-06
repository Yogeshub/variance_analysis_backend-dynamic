from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.platypus import Table
import os

# def generate_pdf(session_id, data):
#
#     file_path = f"temp/{session_id}_report.pdf"
#     doc = SimpleDocTemplate(file_path)
#     elements = []
#     styles = getSampleStyleSheet()
#
#     elements.append(Paragraph("Variance Analysis Report", styles["Title"]))
#     elements.append(Spacer(1, 20))
#
#     for kpi, details in data.items():
#         elements.append(Paragraph(f"<b>{kpi}</b>", styles["Heading2"]))
#         elements.append(Spacer(1, 10))
#
#         summary = details["summary"]
#
#         table_data = [
#             ["Total", summary["total"]],
#             ["Mean", summary["mean"]],
#             ["Std", summary["std"]]
#         ]
#
#         table = Table(table_data)
#         elements.append(table)
#         elements.append(Spacer(1, 20))
#
#     doc.build(elements)
#
#     return file_path

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
import os


def generate_pdf(file_path, analysis_data, ai_data):

    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc = SimpleDocTemplate(file_path)
    elements = []
    styles = getSampleStyleSheet()

    elements.append(Paragraph("KPI Performance Report", styles["Heading1"]))
    elements.append(Spacer(1, 0.3 * inch))

    for kpi, data in analysis_data.items():
        elements.append(Paragraph(f"KPI: {kpi}", styles["Heading2"]))
        elements.append(Spacer(1, 0.2 * inch))

        summary_table = Table(
            [
                ["Total", data["summary"]["total"]],
                ["Mean", data["summary"]["mean"]],
                ["Std", data["summary"]["std"]],
                ["Latest Variance", data["variance_latest"]],
            ]
        )

        summary_table.setStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ]
        )

        elements.append(summary_table)
        elements.append(Spacer(1, 0.3 * inch))

    elements.append(Paragraph("AI Insights", styles["Heading2"]))
    elements.append(Paragraph(ai_data.get("executive_summary", ""), styles["Normal"]))

    doc.build(elements)


from docx import Document
import os

# def generate_docx(session_id, data):
#
#     file_path = f"temp/{session_id}_report.docx"
#     doc = Document()
#     doc.add_heading("Variance Analysis Report", level=1)
#
#     for kpi, details in data.items():
#         doc.add_heading(kpi, level=2)
#         summary = details["summary"]
#
#         doc.add_paragraph(f"Total: {summary['total']}")
#         doc.add_paragraph(f"Mean: {summary['mean']}")
#         doc.add_paragraph(f"Std: {summary['std']}")
#
#     doc.save(file_path)
#
#     return file_path


from docx import Document


def generate_docx(file_path, analysis_data, ai_data):

    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc = Document()
    doc.add_heading("KPI Performance Report", level=1)

    for kpi, data in analysis_data.items():
        doc.add_heading(f"KPI: {kpi}", level=2)

        doc.add_paragraph(f"Total: {data['summary']['total']}")
        doc.add_paragraph(f"Mean: {data['summary']['mean']}")
        doc.add_paragraph(f"Std: {data['summary']['std']}")
        doc.add_paragraph(f"Latest Variance: {data['variance_latest']}")

    doc.add_heading("AI Insights", level=2)
    doc.add_paragraph(ai_data.get("executive_summary", ""))

    doc.save(file_path)
